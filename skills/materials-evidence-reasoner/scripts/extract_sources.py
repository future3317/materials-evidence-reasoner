#!/usr/bin/env python3
"""Extract research sources into LLM-ready, provenance-preserving artifacts.

The command deliberately separates document conversion from scientific
interpretation.  It uses Docling for layout-aware PDFs when available, then
falls back to PyMuPDF/pdfplumber for text-first extraction.  XML/JATS, HTML,
DOCX and XLSX use established public parsers and keep section/table/figure
anchors so an Agent can cite the original source instead of treating converted
text as a new source.

Examples:
    python scripts/extract_sources.py paper.pdf --output literature-extraction
    python scripts/extract_sources.py paper.xml supplement.docx --output extraction
    python scripts/extract_sources.py --check-environment
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import io
import importlib.metadata
import json
import os
import platform
import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

ROOT = Path(__file__).resolve().parents[1]
SCHEMA_PATH = ROOT / "references" / "source-extraction-schema.json"
EXTRACTOR_VERSION = "0.3.1"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def relative_artifact_path(path: Path, base: Path) -> str:
    """Return a portable path relative to the generated bundle root."""
    try:
        return Path(os.path.relpath(path.resolve(), base.resolve())).as_posix()
    except ValueError:
        # Different Windows drives cannot be represented by a relative path;
        # keep the artifact portable rather than leaking an absolute path.
        return path.name


def is_absolute_reference(value: Any) -> bool:
    text = str(value or "")
    # Artifact paths are portable relative paths. Reject drive-relative paths
    # and URI schemes as well as ordinary absolute Windows/POSIX paths.
    return (
        Path(text).is_absolute()
        or bool(re.match(r"^[A-Za-z]:", text))
        or text.startswith("\\\\")
        or bool(re.match(r"^[A-Za-z][A-Za-z0-9+.-]*:", text))
    )


def detected_kind(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        head = path.read_bytes()[:512]
    except OSError:
        head = b""
    # Prefer reliable magic/container signatures when an exported file has
    # the wrong extension. Keep extension-based detection for ordinary text.
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"PK"):
        try:
            import zipfile

            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
            if "[Content_Types].xml" in names:
                if any(name.startswith("word/") for name in names):
                    return "docx"
                if any(name.startswith("xl/") for name in names):
                    return "xlsx"
        except (OSError, ValueError, zipfile.BadZipFile):
            pass
    stripped = head.lstrip().lower()
    if stripped.startswith(b"<?xml") or b"<article" in stripped[:256] or b"<tei" in stripped[:256]:
        return "xml"
    if b"<html" in stripped[:256] or b"<!doctype html" in stripped[:256]:
        return "html"
    if ext == ".pdf":
        return "pdf"
    if ext in {".xml", ".nxml", ".jats", ".tei"}:
        return "xml"
    if ext in {".html", ".htm", ".xhtml"}:
        return "html"
    if ext == ".docx":
        return "docx"
    if ext == ".xlsx":
        return "xlsx"
    if ext == ".csv":
        return "csv"
    if ext == ".tsv":
        return "tsv"
    if ext == ".json":
        return "json"
    if ext in {".md", ".markdown", ".txt", ".rst", ".log", ".py"}:
        return "text"
    return "unsupported"


def iter_inputs(inputs: Iterable[Path], output: Path) -> list[Path]:
    paths: list[Path] = []
    output_resolved = output.resolve()
    for root in inputs:
        root = root.resolve()
        if root.is_file():
            if not root.name.startswith("._") and root.name.lower() not in {".ds_store", "thumbs.db"} and not any(part.casefold() == "__macosx" for part in root.parts):
                paths.append(root)
            continue
        if not root.is_dir():
            raise FileNotFoundError(f"input path does not exist: {root}")
        for path in sorted(root.rglob("*")):
            if not path.is_file():
                continue
            if path.name.startswith("._") or path.name.lower() in {".ds_store", "thumbs.db"} or any(part.casefold() == "__macosx" for part in path.parts):
                continue
            try:
                path.relative_to(output_resolved)
            except ValueError:
                paths.append(path)
    return list(dict.fromkeys(paths))


def package_version(distribution: str) -> str | None:
    try:
        return importlib.metadata.version(distribution)
    except importlib.metadata.PackageNotFoundError:
        return None


def module_available(module: str) -> bool:
    try:
        from importlib.util import find_spec

        return find_spec(module) is not None
    except (ImportError, ModuleNotFoundError, ValueError):
        return False


def docling_model_cache(model_cache: Path | None = None) -> Path:
    """Resolve Docling's model cache without using a user artifact path."""
    if model_cache is not None:
        return model_cache.expanduser().resolve()
    try:
        from docling.datamodel.settings import settings

        return (Path(settings.cache_dir).expanduser() / "models").resolve()
    except (ImportError, AttributeError, TypeError):
        return (Path.home() / ".cache" / "docling" / "models").resolve()


def _has_model_files(path: Path) -> bool:
    return path.is_dir() and any(item.is_file() for item in path.rglob("*"))


def docling_model_status(model_cache: Path | None = None, *, with_ocr: bool = False) -> dict[str, Any]:
    """Describe only the models required by this extractor's PDF profile."""
    cache = docling_model_cache(model_cache)
    result: dict[str, Any] = {"cache_dir": str(cache), "models": [], "missing": [], "ready": False}
    if not module_available("docling"):
        result["status"] = "unavailable"
        result["missing"] = ["docling"]
        return result
    try:
        from docling.datamodel.pipeline_options import LayoutOptions
        from docling.models.stages.layout.layout_model import LayoutModel
        from docling.models.stages.table_structure.table_structure_model import TableStructureModel

        required = [
            ("layout", cache / LayoutOptions().model_spec.model_repo_folder, "布局模型"),
            ("table-structure", cache / TableStructureModel._model_repo_folder, "表格结构模型"),
        ]
        if with_ocr and module_available("rapidocr") and module_available("onnxruntime"):
            from docling.models.stages.ocr.rapid_ocr_model import RapidOcrModel

            required.append(("ocr-rapidocr", cache / RapidOcrModel._model_repo_folder, "RapidOCR 中文模型"))
        for model_id, path, label in required:
            present = _has_model_files(path)
            item = {"id": model_id, "label": label, "path": str(path), "present": present}
            result["models"].append(item)
            if not present:
                result["missing"].append(model_id)
        result["ready"] = not result["missing"]
        result["status"] = "ready" if result["ready"] else "missing"
        if with_ocr and not any(item["id"] == "ocr-rapidocr" for item in result["models"]):
            result["ocr_note"] = "RapidOCR/onnxruntime 未同时可用；Docling 将按 Auto OCR 运行时选择其他引擎，或明确记录没有 OCR 引擎。"
        return result
    except Exception as exc:
        result["status"] = "inspection-failed"
        result["error"] = f"{type(exc).__name__}: {exc}"
        return result


def ensure_docling_models(
    model_cache: Path | None = None,
    *,
    with_ocr: bool = False,
    allow_download: bool = True,
    force: bool = False,
) -> dict[str, Any]:
    """Prepare required Docling models once per extraction run.

    Downloads use Docling's public model APIs and are intentionally limited to
    layout, table structure, and the selected RapidOCR backend. Optional VLM,
    formula, picture-classifier, and alternate OCR models are not downloaded.
    """
    status = docling_model_status(model_cache, with_ocr=with_ocr)
    status.update({"download_requested": allow_download, "download_performed": False, "force": force})
    if status.get("status") == "unavailable":
        return status
    if status.get("status") == "ready" and not force:
        return status
    if not status.get("missing") and not force:
        return status
    if not allow_download:
        status["status"] = "skipped"
        status["next_step"] = "移除 --no-model-download，或先手动准备 model-cache 后重试。"
        return status
    cache = docling_model_cache(model_cache)
    try:
        from docling.datamodel.pipeline_options import LayoutOptions
        from docling.models.stages.layout.layout_model import LayoutModel
        from docling.models.stages.table_structure.table_structure_model import TableStructureModel

        cache.mkdir(parents=True, exist_ok=True)
        if force or "layout" in status["missing"]:
            LayoutModel.download_models(local_dir=cache / LayoutOptions().model_spec.model_repo_folder, force=force, progress=True)
        if force or "table-structure" in status["missing"]:
            TableStructureModel.download_models(local_dir=cache / TableStructureModel._model_repo_folder, force=force, progress=True)
        if force or "ocr-rapidocr" in status["missing"]:
            from docling.models.stages.ocr.rapid_ocr_model import RapidOcrModel

            RapidOcrModel.download_models(backend="onnxruntime", local_dir=cache / RapidOcrModel._model_repo_folder, force=force, progress=True, lang="chinese")
        status = docling_model_status(cache, with_ocr=with_ocr)
        status.update({"download_requested": allow_download, "download_performed": True, "force": force})
        if not status.get("ready"):
            status["status"] = "incomplete"
            status["next_step"] = "模型下载后仍有缺失；检查网络、磁盘空间和缓存目录。"
        return status
    except Exception as exc:
        status.update({"status": "download-failed", "download_performed": True, "error": f"{type(exc).__name__}: {exc}", "next_step": "检查网络、磁盘空间和 Hugging Face 可访问性；auto PDF 会尝试文本后端回退。"})
        return status


def environment_report(model_cache: Path | None = None, *, with_ocr: bool = False) -> dict[str, Any]:
    packages = {
        "docling": package_version("docling"),
        "PyMuPDF": package_version("PyMuPDF"),
        "pdfplumber": package_version("pdfplumber"),
        "lxml": package_version("lxml"),
        "beautifulsoup4": package_version("beautifulsoup4"),
        "python-docx": package_version("python-docx"),
        "openpyxl": package_version("openpyxl"),
        "defusedxml": package_version("defusedxml"),
        "jsonschema": package_version("jsonschema"),
    }
    modules = {
        "docling": module_available("docling"),
        "fitz": module_available("fitz"),
        "pdfplumber": module_available("pdfplumber"),
        "lxml": module_available("lxml"),
        "bs4": module_available("bs4"),
        "docx": module_available("docx"),
        "openpyxl": module_available("openpyxl"),
        "defusedxml": module_available("defusedxml"),
        "jsonschema": module_available("jsonschema"),
    }
    tools = {name: shutil.which(name) for name in ("pdftotext", "pdftoppm", "tesseract")}
    capabilities = {
        "pdf_layout": modules["docling"],
        "pdf_text": modules["fitz"] or modules["pdfplumber"],
        "pdf_tables": modules["docling"] or modules["pdfplumber"],
        "pdf_page_rendering": modules["fitz"] or bool(tools["pdftoppm"]),
        "pdf_ocr": modules["docling"],
        "xml_jats": modules["lxml"],
        "html": modules["bs4"],
        "docx": modules["docx"],
        "xlsx": modules["openpyxl"],
        "schema_validation": modules["jsonschema"],
    }
    missing: list[str] = []
    if not modules["docling"] and not modules["fitz"] and not modules["pdfplumber"]:
        missing.append("PDF extraction: install docling or PyMuPDF/pdfplumber")
    if not modules["lxml"]:
        missing.append("XML extraction: install lxml")
    if not modules["bs4"]:
        missing.append("HTML extraction: install beautifulsoup4")
    if not modules["docx"]:
        missing.append("DOCX extraction: install python-docx")
    if not modules["openpyxl"]:
        missing.append("XLSX extraction: install openpyxl")
    profiles = {
        "pdf_auto": {
            "status": "ready" if any(modules[name] for name in ("docling", "fitz", "pdfplumber")) else "unavailable",
            "backends": [name for name, available in (("docling", modules["docling"]), ("pymupdf", modules["fitz"]), ("pdfplumber", modules["pdfplumber"])) if available],
            "next_step": "直接运行 extract_sources.py；首次使用会自动尝试下载必要 Docling 模型。" if modules["docling"] else "安装 docling，或显式使用 PyMuPDF/pdfplumber。",
        },
        "pdf_ocr": {
            "status": "available-needs-runtime-check" if modules["docling"] else "unavailable",
            "backend": "docling" if modules["docling"] else None,
            "next_step": "对扫描件使用 --ocr；若首次运行失败，先用 --render-pages 检查页图并记录限制。" if modules["docling"] else "安装 docling；PyMuPDF/pdfplumber 不执行 OCR。",
        },
        "xml": {"status": "ready" if modules["lxml"] else "unavailable", "next_step": "使用 lxml 安全解析。" if modules["lxml"] else "安装 lxml。"},
        "html": {"status": "ready" if modules["bs4"] else "unavailable", "next_step": "使用 Beautiful Soup 解析。" if modules["bs4"] else "安装 beautifulsoup4。"},
        "docx": {"status": "ready" if modules["docx"] else "unavailable", "next_step": "使用 python-docx 解析正文、表格和可选图像。" if modules["docx"] else "安装 python-docx。"},
        "xlsx": {"status": "ready" if modules["openpyxl"] else "unavailable", "next_step": "使用 openpyxl；公式只保留公式与缓存值，不执行公式。" if modules["openpyxl"] else "安装 openpyxl。"},
        "bundle_validation": {"status": "ready" if modules["jsonschema"] else "structural-only", "next_step": "使用 jsonschema 做完整校验。" if modules["jsonschema"] else "安装 jsonschema 以启用完整校验。"},
    }
    actions = [
        {"id": "pdf-first-pass", "when": "PDF 有文本且不需要 OCR", "command": "python scripts/extract_sources.py paper.pdf --output source-extraction", "reason": "先用自动后端获取带锚点文本和表格。"},
        {"id": "pdf-scanned", "when": "PDF 是扫描件或提取为空", "command": "python scripts/extract_sources.py paper.pdf --ocr --render-pages --output source-extraction", "reason": "让 Docling 尝试 OCR，并保留页图供人工复核。"},
        {"id": "pdf-fast-offline", "when": "需要避免布局模型或快速预览", "command": "python scripts/extract_sources.py paper.pdf --pdf-backend pymupdf --output source-extraction", "reason": "文本优先回退；结果仍需检查双栏、表格和图注。"},
    ]
    return {
        "extractor_version": EXTRACTOR_VERSION,
        "python": sys.version,
        "python_executable": sys.executable,
        "platform": platform.platform(),
        "packages": packages,
        "modules": modules,
        "external_tools": tools,
        "capabilities": capabilities,
        "install_hints": {
            "pdf_layout": "pip/conda install docling",
            "pdf_text": "pip/conda install PyMuPDF pdfplumber",
            "xml_jats": "pip/conda install lxml",
            "html": "pip/conda install beautifulsoup4",
            "docx": "pip/conda install python-docx",
            "xlsx": "pip/conda install openpyxl",
            "schema_validation": "pip/conda install jsonschema",
        },
        "profiles": {
            "pdf_auto": "docling -> pymupdf -> pdfplumber",
            "pdf_offline": "pymupdf -> pdfplumber",
            "xml": "lxml safe parser (defusedxml optional)",
            "html": "beautifulsoup4",
            "docx": "python-docx",
            "xlsx": "openpyxl",
            "bundle_validation": "jsonschema (optional; structural fallback otherwise)",
        },
        "profile_status": profiles,
        "recommended_actions": actions,
        "model_cache": docling_model_status(model_cache, with_ocr=with_ocr),
        "warnings": [
            "首次使用 Docling PDF profile 会自动尝试下载必要模型；--no-model-download 或 --pdf-backend pymupdf 可用于离线运行。",
            "Converted text is an extraction artifact, not a scientific conclusion or an independent source.",
        ],
        "missing": missing,
    }


def compact_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def local_name(tag: Any) -> str:
    if not isinstance(tag, str):
        return ""
    return tag.rsplit("}", 1)[-1].lower()


def xpath_local(element: Any, names: Iterable[str]) -> list[Any]:
    names = list(names)
    expression = ".//*[" + " or ".join(f"local-name()='{name}'" for name in names) + "]"
    return list(element.xpath(expression))


def first_text(element: Any, names: Iterable[str]) -> str:
    for child in xpath_local(element, names):
        value = compact_text(" ".join(child.itertext()))
        if value:
            return value
    return ""


def anchor(source_id: str, *, page: int | None = None, item: str | None = None, section: str | None = None) -> str:
    parts = [source_id]
    if page is not None:
        parts.append(f"page={page}")
    if item:
        parts.append(f"item={item}")
    if section:
        parts.append(f"section={section}")
    return ":".join(parts)


def empty_document(source_id: str, path: Path, kind: str) -> dict[str, Any]:
    return {
        "source_id": source_id,
        "name": path.name,
        "kind": kind,
        "title": path.stem,
        "metadata": {},
        "sections": [],
        "blocks": [],
        "tables": [],
        "figures": [],
        "references": [],
        "reading_order": [],
        "warnings": [],
    }


def table_markdown(headers: list[str], rows: list[list[str]]) -> str:
    if not headers and not rows:
        return ""
    width = max(len(headers), max((len(row) for row in rows), default=0), 1)
    headers = (headers + [""] * width)[:width]
    def display(value: Any) -> str:
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False, separators=(",", ":"))
        return str(value if value is not None else "")

    lines = ["| " + " | ".join(display(item).replace("|", "\\|") for item in headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    for row in rows:
        values = (row + [""] * width)[:width]
        lines.append("| " + " | ".join(display(item).replace("|", "\\|") for item in values) + " |")
    return "\n".join(lines)


def write_table_csv(table: dict[str, Any], output: Path) -> str:
    output.parent.mkdir(parents=True, exist_ok=True)
    headers = list(table.get("headers", []))
    rows = [list(row) for row in table.get("rows", [])]
    width = max(len(headers), max((len(row) for row in rows), default=0), 1)
    headers = (headers + [f"column_{i}" for i in range(len(headers) + 1, width + 1)])[:width]
    with output.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(headers)
        for row in rows:
            values = (row + [""] * width)[:width]
            writer.writerow(json.dumps(value, ensure_ascii=False, separators=(",", ":")) if isinstance(value, (dict, list)) else value for value in values)
    return str(output)


def write_markdown(document: dict[str, Any], output: Path) -> None:
    source_id = document["source_id"]
    lines = [f"# {document.get('title') or document['name']}", "", f"<!-- extracted-source: {source_id}; kind={document['kind']} -->", ""]
    metadata = document.get("metadata") or {}
    for key, value in metadata.items():
        if value not in (None, "", []):
            lines.append(f"- {key}: {value}")
    if metadata:
        lines.append("")
    order = document.get("reading_order") or []
    if order:
        sections = {item.get("anchor"): item for item in document.get("sections", [])}
        blocks = {item.get("anchor"): item for item in document.get("blocks", [])}
        tables = {item.get("id"): item for item in document.get("tables", [])}
        figures = {item.get("id"): item for item in document.get("figures", [])}
        for item in order:
            kind = item.get("type")
            ref = item.get("ref")
            if kind == "section" and ref in sections:
                section = sections[ref]
                level = max(2, min(6, int(section.get("level", 2))))
                lines.extend([f"{'#' * level} {section.get('heading') or '未命名章节'}", "", f"<!-- source: {section['anchor']} -->", ""])
            elif kind == "block" and ref in blocks:
                block = blocks[ref]
                if block.get("text"):
                    context = f"; section={block['section_anchor']}" if block.get("section_anchor") else ""
                    lines.extend([f"<!-- source: {block['anchor']}{context} -->", block["text"], ""])
            elif kind == "table" and ref in tables:
                table = tables[ref]
                lines.extend([f"### {table.get('caption') or table['id']}", "", f"<!-- source: {table['anchor']} -->", table_markdown(table.get("headers", []), table.get("rows", [])), ""])
            elif kind == "figure" and ref in figures:
                figure = figures[ref]
                lines.extend([f"### Figure {figure['id']}", "", f"<!-- source: {figure['anchor']} -->", figure.get("caption") or "图注未提取。", ""])
                if figure.get("image_path"):
                    lines.extend([f"![{figure.get('caption') or figure['id']}](../{figure['image_path']})", ""])
    else:
        for section in document.get("sections", []):
            level = max(2, min(6, int(section.get("level", 2))))
            lines.extend([f"{'#' * level} {section.get('heading') or '未命名章节'}", "", f"<!-- source: {section['anchor']} -->", section.get("text", ""), ""])
        for block in document.get("blocks", []):
            if block.get("text"):
                lines.extend([f"<!-- source: {block['anchor']} -->", block["text"], ""])
        for table in document.get("tables", []):
            lines.extend([f"### {table.get('caption') or table['id']}", "", f"<!-- source: {table['anchor']} -->", table_markdown(table.get("headers", []), table.get("rows", [])), ""])
        for figure in document.get("figures", []):
            lines.extend([f"### Figure {figure['id']}", "", f"<!-- source: {figure['anchor']} -->", figure.get("caption") or "图注未提取。", ""])
            if figure.get("image_path"):
                lines.extend([f"![{figure.get('caption') or figure['id']}](../{figure['image_path']})", ""])
    if document.get("references"):
        lines.extend(["## References", ""])
        lines.extend(f"- [{item['id']}] {item.get('text', '')}" for item in document["references"])
        lines.append("")
    if document.get("warnings"):
        lines.extend(["## Extraction notes", ""])
        lines.extend(f"- {warning}" for warning in document["warnings"])
        lines.append("")
    output.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def extract_xml(path: Path, source_id: str) -> dict[str, Any]:
    document = empty_document(source_id, path, "xml")
    document["extractor"] = "lxml-safe-parser"
    try:
        from lxml import etree

        parser = etree.XMLParser(resolve_entities=False, no_network=True, recover=False, huge_tree=True)
        root = etree.parse(str(path), parser).getroot()
    except Exception as exc:
        document["warnings"].append(f"XML parse failed: {type(exc).__name__}: {exc}")
        return document

    document["title"] = first_text(root, ("article-title", "title")) or path.stem
    document["metadata"] = {
        "doi": next((compact_text(node.text) for node in xpath_local(root, ("article-id",)) if node.get("pub-id-type") == "doi" and compact_text(node.text)), None),
        "authors": ", ".join(compact_text(" ".join(node.itertext())) for node in xpath_local(root, ("name",)) if compact_text(" ".join(node.itertext())))[:2000],
    }
    for abstract_index, node in enumerate(xpath_local(root, ("abstract",)), start=1):
        text = compact_text(" ".join(node.itertext()))
        if text:
            document["sections"].append({"anchor": anchor(source_id, section=f"abstract-{abstract_index}"), "heading": "Abstract", "level": 2, "text": text})
    for index, node in enumerate(xpath_local(root, ("sec",)), start=1):
        title = first_text(node, ("title",))
        paragraphs: list[str] = []
        for paragraph in xpath_local(node, ("p",)):
            nearest_sec = next((parent for parent in paragraph.iterancestors() if local_name(parent) == "sec"), None)
            if nearest_sec is node:
                value = compact_text(" ".join(paragraph.itertext()))
                if value:
                    paragraphs.append(value)
        if not paragraphs:
            value = compact_text(" ".join(node.itertext()))
            if title and value.startswith(title):
                value = value[len(title):].strip()
            paragraphs = [value] if value else []
        if paragraphs:
            document["sections"].append({"anchor": anchor(source_id, section=node.get("id") or f"sec-{index}"), "heading": title or f"Section {index}", "level": 2, "text": "\n\n".join(paragraphs)})
    table_nodes = xpath_local(root, ("table-wrap",))
    if not table_nodes:
        table_nodes = xpath_local(root, ("table",))
    for index, wrapper in enumerate(table_nodes, start=1):
        table = next(iter(xpath_local(wrapper, ("table",))), wrapper)
        rows = []
        for row in xpath_local(table, ("tr",)):
            cells = xpath_local(row, ("th", "td"))
            if cells:
                rows.append([compact_text(" ".join(cell.itertext())) for cell in cells])
        if rows:
            caption = first_text(wrapper, ("caption", "title", "label"))
            document["tables"].append({"id": f"{source_id}-T{index:03d}", "anchor": anchor(source_id, section=wrapper.get("id") or f"table-{index}"), "caption": caption, "headers": rows[0], "rows": rows[1:], "page": None})
    for index, figure in enumerate(xpath_local(root, ("fig", "figure")), start=1):
        caption = first_text(figure, ("caption", "title", "label"))
        graphic = next(iter(xpath_local(figure, ("graphic", "inline-graphic"))), None)
        href = ""
        if graphic is not None:
            href = next((value for key, value in graphic.attrib.items() if local_name(key) == "href"), "")
        document["figures"].append({"id": f"{source_id}-F{index:03d}", "anchor": anchor(source_id, section=figure.get("id") or f"figure-{index}"), "caption": caption, "page": None, "source_ref": href})
    for index, reference in enumerate(xpath_local(root, ("ref",)), start=1):
        value = compact_text(" ".join(reference.itertext()))
        if value:
            document["references"].append({"id": reference.get("id") or f"{source_id}-R{index:03d}", "text": value})
    return document


def extract_html(path: Path, source_id: str) -> dict[str, Any]:
    document = empty_document(source_id, path, "html")
    document["extractor"] = "beautifulsoup4"
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
    except Exception as exc:
        document["warnings"].append(f"HTML parse failed: {type(exc).__name__}: {exc}")
        return document
    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    document["title"] = title or path.stem
    for index, heading in enumerate(soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]), start=1):
        text = heading.get_text(" ", strip=True)
        parts: list[str] = []
        for sibling in heading.find_all_next():
            if sibling.name in {"h1", "h2", "h3", "h4", "h5", "h6"} and sibling is not heading:
                break
            if sibling.name in {"p", "li"}:
                value = sibling.get_text(" ", strip=True)
                if value and value not in parts:
                    parts.append(value)
        document["sections"].append({"anchor": anchor(source_id, section=f"html-{index}"), "heading": text or f"Section {index}", "level": int(heading.name[1]) + 1, "text": "\n\n".join(parts)})
    if not document["sections"]:
        text = soup.get_text("\n", strip=True)
        document["blocks"].append({"anchor": anchor(source_id, section="body"), "page": None, "label": "text", "text": text})
    for index, table in enumerate(soup.find_all("table"), start=1):
        rows = [[compact_text(cell.get_text(" ", strip=True)) for cell in row.find_all(["th", "td"])] for row in table.find_all("tr")]
        rows = [row for row in rows if row]
        if rows:
            document["tables"].append({"id": f"{source_id}-T{index:03d}", "anchor": anchor(source_id, section=f"table-{index}"), "caption": "", "headers": rows[0], "rows": rows[1:], "page": None})
    return document


def source_text(path: Path) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding), encoding
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace"), "utf-8-replace"


def extract_delimited(path: Path, source_id: str) -> dict[str, Any]:
    kind = "tsv" if path.suffix.lower() == ".tsv" else "csv"
    document = empty_document(source_id, path, kind)
    document["extractor"] = "python-csv"
    try:
        text, encoding = source_text(path)
        sample = text[:8192]
        delimiter = "\t" if kind == "tsv" else ","
        try:
            delimiter = csv.Sniffer().sniff(sample, delimiters=",\t;|").delimiter
        except csv.Error:
            pass
        # Feed the complete text stream to csv.reader so quoted fields may
        # legitimately contain newlines without being split into fake rows.
        rows = [row for row in csv.reader(io.StringIO(text), delimiter=delimiter) if any(str(value).strip() for value in row)]
    except Exception as exc:
        document["warnings"].append(f"delimited text parse failed: {type(exc).__name__}: {exc}")
        return document
    document["metadata"] = {"encoding": encoding, "delimiter": delimiter}
    if rows:
        document["tables"].append({"id": f"{source_id}-T001", "anchor": anchor(source_id, section="table-1"), "caption": path.name, "headers": rows[0], "rows": rows[1:], "page": None})
    else:
        document["warnings"].append("delimited file has no non-empty rows")
    return document


def extract_json(path: Path, source_id: str) -> dict[str, Any]:
    document = empty_document(source_id, path, "json")
    document["extractor"] = "python-json"
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        document["warnings"].append(f"JSON parse failed: {type(exc).__name__}: {exc}")
        return document
    document["title"] = path.stem
    arrays: list[tuple[str, list[dict[str, Any]]]] = []
    if isinstance(data, list) and all(isinstance(item, dict) for item in data):
        arrays.append(("root", data))
    elif isinstance(data, dict):
        arrays = [(str(key), value) for key, value in data.items() if isinstance(value, list) and value and all(isinstance(item, dict) for item in value)]
        scalar = {key: value for key, value in data.items() if not isinstance(value, (dict, list))}
        if scalar:
            document["metadata"]["top_level_scalars"] = scalar
    if not arrays:
        document["blocks"].append({"anchor": anchor(source_id, section="body"), "page": None, "label": "json", "text": json.dumps(data, ensure_ascii=False, indent=2)})
        return document
    document["metadata"].update({"arrays": [name for name, _ in arrays], "array_count": len(arrays)})
    for table_index, (name, items) in enumerate(arrays, start=1):
        headers: list[str] = []
        for item in items:
            for key in item:
                if key not in headers:
                    headers.append(key)
        rows = [[item.get(key) for key in headers] for item in items]
        table_id = f"{source_id}-T{table_index:03d}"
        document["tables"].append({"id": table_id, "anchor": anchor(source_id, section=f"array-{name}"), "caption": f"{path.name} · {name}", "headers": headers, "rows": rows, "page": None})
        document["reading_order"].append({"type": "table", "ref": table_id})
    return document


def extract_docx(path: Path, source_id: str, extract_figures: bool = False) -> dict[str, Any]:
    document = empty_document(source_id, path, "docx")
    document["extractor"] = "python-docx"
    try:
        from docx import Document

        source = Document(str(path))
    except Exception as exc:
        document["warnings"].append(f"DOCX parse failed: {type(exc).__name__}: {exc}")
        return document
    document["title"] = path.stem
    if extract_figures:
        document["_figure_images"] = {}
    section_index = 0
    for paragraph in source.paragraphs:
        text = compact_text(paragraph.text)
        style = str(getattr(paragraph.style, "name", ""))
        if text and style.lower().startswith("heading"):
            match = re.search(r"(\d+)", style)
            level = int(match.group(1)) + 1 if match else 2
            section_index += 1
            section_anchor = anchor(source_id, section=f"docx-{section_index}")
            document["sections"].append({"anchor": section_anchor, "heading": text, "level": level, "text": ""})
            document["reading_order"].append({"type": "section", "ref": section_anchor})
        elif text and document["sections"]:
            document["sections"][-1]["text"] = (document["sections"][-1]["text"] + "\n\n" + text).strip()
        elif text:
            block_anchor = anchor(source_id, section=f"docx-body-{len(document['blocks']) + 1}")
            document["blocks"].append({"anchor": block_anchor, "page": None, "label": "paragraph", "text": text})
            document["reading_order"].append({"type": "block", "ref": block_anchor})
        if extract_figures:
            for element in paragraph._p.iter():
                if local_name(getattr(element, "tag", "")) != "blip":
                    continue
                relationship_id = element.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                image_part = getattr(source.part, "related_parts", {}).get(relationship_id)
                blob = getattr(image_part, "blob", None)
                if not blob:
                    document["warnings"].append(f"DOCX embedded image could not be read: {relationship_id or 'unknown relationship'}")
                    continue
                figure_id = f"{source_id}-F{len(document['figures']) + 1:03d}"
                document["figures"].append({"id": figure_id, "anchor": anchor(source_id, section=f"figure-{len(document['figures']) + 1}"), "caption": "", "page": None})
                document["reading_order"].append({"type": "figure", "ref": figure_id})
                document["_figure_images"][figure_id] = image_part
    for index, table in enumerate(source.tables, start=1):
        rows = [[compact_text(cell.text) for cell in row.cells] for row in table.rows]
        if rows:
            table_id = f"{source_id}-T{index:03d}"
            document["tables"].append({"id": table_id, "anchor": anchor(source_id, section=f"table-{index}"), "caption": "", "headers": rows[0], "rows": rows[1:], "page": None})
            document["reading_order"].append({"type": "table", "ref": table_id})
    if extract_figures and not document.get("_figure_images"):
        document.pop("_figure_images", None)
    return document


def extract_xlsx(path: Path, source_id: str) -> dict[str, Any]:
    document = empty_document(source_id, path, "xlsx")
    document["extractor"] = "openpyxl"
    try:
        import openpyxl

        # Keep the formula workbook normal-sized so hidden/merged metadata is
        # available; the cached-value view remains streaming-friendly.
        workbook = openpyxl.load_workbook(path, data_only=False, read_only=False)
        cached_workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception as exc:
        document["warnings"].append(f"XLSX parse failed: {type(exc).__name__}: {exc}")
        return document
    document["title"] = path.stem
    for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
        cached_sheet = cached_workbook[sheet.title]
        rows: list[list[Any]] = []
        formula_count = 0
        missing_cache_count = 0
        for formula_row, cached_row in zip(sheet.iter_rows(), cached_sheet.iter_rows()):
            output_row: list[Any] = []
            for formula_cell, cached_cell in zip(formula_row, cached_row):
                value = formula_cell.value
                if isinstance(value, str) and value.startswith("="):
                    formula_count += 1
                    cached_value = cached_cell.value
                    if cached_value is None:
                        missing_cache_count += 1
                    output_row.append({"formula": value, "cached_value": cached_value})
                else:
                    output_row.append(value)
            rows.append(output_row)
        rows = [row for row in rows if any(value is not None and (bool(value) if isinstance(value, (dict, list)) else bool(str(value).strip())) for value in row)]
        if not rows:
            continue
        warnings: list[str] = []
        if formula_count:
            warnings.append(f"contains {formula_count} formula cells; formulas were not executed")
        if missing_cache_count:
            warnings.append(f"{missing_cache_count} formula cells have no cached value")
        hidden_rows = sum(1 for dimension in sheet.row_dimensions.values() if dimension.hidden)
        hidden_cols = sum(1 for dimension in sheet.column_dimensions.values() if dimension.hidden)
        if hidden_rows or hidden_cols:
            warnings.append(f"hidden rows={hidden_rows}, hidden columns={hidden_cols}")
        merged = len(sheet.merged_cells.ranges)
        if merged:
            warnings.append(f"worksheet has {merged} merged ranges")
        document["tables"].append({"id": f"{source_id}-T{sheet_index:03d}", "anchor": anchor(source_id, section=f"sheet-{sheet.title}"), "caption": sheet.title, "headers": rows[0], "rows": rows[1:], "page": None, "sheet_name": sheet.title, "warnings": warnings})
    workbook.close()
    cached_workbook.close()
    return document


def item_label(item: Any) -> str:
    return str(getattr(getattr(item, "label", None), "value", getattr(item, "label", "item")))


def item_pages(item: Any) -> list[int]:
    pages: set[int] = set()
    for provenance in getattr(item, "prov", []) or []:
        page = getattr(provenance, "page_no", None)
        if isinstance(page, int):
            pages.add(page)
    return sorted(pages)


def extract_pdf_docling(
    path: Path,
    source_id: str,
    no_ocr: bool = True,
    extract_figures: bool = False,
    model_cache: Path | None = None,
) -> dict[str, Any]:
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption

    option_kwargs: dict[str, Any] = {
        "do_ocr": not no_ocr,
        "do_table_structure": True,
        "generate_page_images": False,
        "generate_picture_images": extract_figures,
    }
    if model_cache is not None:
        option_kwargs["artifacts_path"] = model_cache.expanduser().resolve()
    options = PdfPipelineOptions(**option_kwargs)
    converter = DocumentConverter(allowed_formats=[InputFormat.PDF], format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=options)})
    result = converter.convert(path)
    document = empty_document(source_id, path, "pdf")
    document["extractor"] = "docling"
    doc = result.document
    document["title"] = path.stem
    document["page_count"] = doc.num_pages() if callable(getattr(doc, "num_pages", None)) else getattr(doc, "num_pages", None)
    current_section_anchor: str | None = None
    document.setdefault("_figure_images", {})
    for index, (item, level) in enumerate(doc.iterate_items(traverse_pictures=True), start=1):
        text = compact_text(getattr(item, "text", ""))
        label = item_label(item)
        pages = item_pages(item)
        page = pages[0] if pages else None
        item_ref = str(getattr(item, "self_ref", f"item-{index}")).lstrip("#/").replace("/", "-")
        if label == "section_header" and text:
            section_anchor = anchor(source_id, page=page, item=item_ref)
            document["sections"].append({"anchor": section_anchor, "heading": text, "level": int(getattr(item, "level", level) or level) + 1, "text": "", "order": index})
            document["reading_order"].append({"type": "section", "ref": section_anchor})
            current_section_anchor = section_anchor
        elif label == "table":
            try:
                dataframe = item.export_to_dataframe(doc)
                headers = [str(value) for value in dataframe.columns.tolist()]
                rows = [["" if value is None else str(value) for value in row] for row in dataframe.fillna("").values.tolist()]
            except Exception as exc:
                headers, rows = [], []
                document["warnings"].append(f"Docling table export failed at {item_ref}: {exc}")
            table_id = f"{source_id}-T{len(document['tables']) + 1:03d}"
            document["tables"].append({"id": table_id, "anchor": anchor(source_id, page=page, item=item_ref), "caption": "", "headers": headers, "rows": rows, "page": page, "order": index, "section_anchor": current_section_anchor})
            document["reading_order"].append({"type": "table", "ref": table_id})
        elif label in {"picture", "figure"}:
            figure_id = f"{source_id}-F{len(document['figures']) + 1:03d}"
            figure = {"id": figure_id, "anchor": anchor(source_id, page=page, item=item_ref), "caption": text, "page": page, "order": index, "section_anchor": current_section_anchor}
            document["figures"].append(figure)
            document["reading_order"].append({"type": "figure", "ref": figure_id})
            if extract_figures:
                try:
                    image = item.get_image(doc)
                    if image is not None:
                        document["_figure_images"][figure_id] = image
                except Exception as exc:
                    document["warnings"].append(f"Docling figure export failed at {item_ref}: {type(exc).__name__}: {exc}")
        elif text:
            block_anchor = anchor(source_id, page=page, item=item_ref)
            document["blocks"].append({"anchor": block_anchor, "page": page, "label": label, "text": text, "order": index, "section_anchor": current_section_anchor})
            document["reading_order"].append({"type": "block", "ref": block_anchor})
    return document


def extract_pdf_pymupdf(path: Path, source_id: str) -> dict[str, Any]:
    import fitz

    document = empty_document(source_id, path, "pdf")
    document["extractor"] = "pymupdf"
    source = fitz.open(str(path))
    document["page_count"] = source.page_count
    for page_index, page in enumerate(source, start=1):
        blocks = page.get_text("blocks")
        for block_index, block in enumerate(sorted(blocks, key=lambda value: (value[1], value[0])), start=1):
            text = compact_text(block[4] if len(block) > 4 else "")
            if text:
                document["blocks"].append({"anchor": anchor(source_id, page=page_index, item=f"block-{block_index}"), "page": page_index, "label": "text", "text": text})
    source.close()
    return document


def extract_pdf_pdfplumber(path: Path, source_id: str) -> dict[str, Any]:
    import pdfplumber

    document = empty_document(source_id, path, "pdf")
    document["extractor"] = "pdfplumber"
    with pdfplumber.open(str(path)) as source:
        document["page_count"] = len(source.pages)
        for page_index, page in enumerate(source.pages, start=1):
            text = compact_text(page.extract_text() or "")
            if text:
                document["blocks"].append({"anchor": anchor(source_id, page=page_index, item="text"), "page": page_index, "label": "text", "text": text})
            for table_index, rows in enumerate(page.extract_tables() or [], start=1):
                clean_rows = [[compact_text(value) for value in row] for row in rows if row]
                if clean_rows:
                    document["tables"].append({"id": f"{source_id}-T{len(document['tables']) + 1:03d}", "anchor": anchor(source_id, page=page_index, item=f"table-{table_index}"), "caption": "", "headers": clean_rows[0], "rows": clean_rows[1:], "page": page_index})
    return document


def extract_pdf(
    path: Path,
    source_id: str,
    backend: str,
    no_ocr: bool = True,
    extract_figures: bool = False,
    model_cache: Path | None = None,
) -> dict[str, Any]:
    choices = [backend] if backend != "auto" else ["docling", "pymupdf", "pdfplumber"]
    errors: list[str] = []
    attempts: list[dict[str, str]] = []
    for choice in choices:
        try:
            if choice == "docling":
                document = extract_pdf_docling(path, source_id, no_ocr, extract_figures, model_cache)
            elif choice == "pymupdf":
                document = extract_pdf_pymupdf(path, source_id)
            elif choice == "pdfplumber":
                document = extract_pdf_pdfplumber(path, source_id)
            else:
                raise ValueError(f"unsupported PDF backend: {choice}")
            if not any(document.get(key) for key in ("blocks", "sections", "tables", "figures", "references")):
                if no_ocr:
                    document["warnings"].append(
                        "no text, tables, figures, or references were extracted; OCR is disabled. "
                        "For a scanned PDF, retry with --ocr and review --render-pages."
                    )
                else:
                    document["warnings"].append(
                        "no text, tables, figures, or references were extracted; review the rendered page and source PDF."
                    )
            if not no_ocr and choice != "docling":
                document["warnings"].append(
                    f"OCR was requested, but the {choice} backend does not perform OCR; use Docling or install it for image-only PDFs."
                )
            document["warnings"].extend(errors)
            document["backend_attempts"] = attempts + [{"backend": choice, "status": "used"}]
            return document
        except ImportError as exc:
            attempts.append({"backend": choice, "status": "unavailable"})
            errors.append(f"{choice} unavailable: {exc}")
        except Exception as exc:
            attempts.append({"backend": choice, "status": "failed"})
            errors.append(f"{choice} failed: {type(exc).__name__}: {exc}")
            if backend != "auto":
                break
    document = empty_document(source_id, path, "pdf")
    document["warnings"].extend(errors or ["no PDF extractor was available"])
    document["backend_attempts"] = attempts
    return document


def extract_one(
    path: Path,
    source_id: str,
    backend: str,
    no_ocr: bool = True,
    extract_figures: bool = False,
    model_cache: Path | None = None,
) -> dict[str, Any]:
    kind = detected_kind(path)
    if kind == "pdf":
        return extract_pdf(path, source_id, backend, no_ocr, extract_figures, model_cache)
    if kind == "xml":
        return extract_xml(path, source_id)
    if kind == "html":
        return extract_html(path, source_id)
    if kind in {"csv", "tsv"}:
        return extract_delimited(path, source_id)
    if kind == "json":
        return extract_json(path, source_id)
    if kind == "docx":
        return extract_docx(path, source_id, extract_figures)
    if kind == "xlsx":
        return extract_xlsx(path, source_id)
    if kind == "text":
        document = empty_document(source_id, path, kind)
        document["extractor"] = "python-text"
        document["title"] = path.stem
        document["blocks"].append({"anchor": anchor(source_id, section="body"), "page": None, "label": "text", "text": path.read_text(encoding="utf-8", errors="replace")})
        return document
    document = empty_document(source_id, path, kind)
    document["warnings"].append("unsupported file type; use prepare_intake.py for manifest registration")
    return document


def render_pdf_pages(path: Path, source_id: str, output: Path, dpi: int) -> list[str]:
    import fitz

    if dpi < 72 or dpi > 300:
        raise ValueError("--render-dpi must be between 72 and 300")
    page_dir = output / "pages"
    page_dir.mkdir(parents=True, exist_ok=True)
    scale = dpi / 72.0
    matrix = fitz.Matrix(scale, scale)
    paths: list[str] = []
    source = fitz.open(str(path))
    for page_index, page in enumerate(source, start=1):
        destination = page_dir / f"{source_id}-P{page_index:04d}.png"
        page.get_pixmap(matrix=matrix, alpha=False).save(str(destination))
        paths.append(destination.relative_to(output).as_posix())
    source.close()
    return paths


def write_document_artifacts(document: dict[str, Any], output: Path) -> dict[str, str]:
    source_id = document["source_id"]
    content_dir = output / "documents"
    table_dir = output / "tables"
    content_dir.mkdir(parents=True, exist_ok=True)
    structure_path = content_dir / f"{source_id}.json"
    markdown_path = content_dir / f"{source_id}.md"
    figure_images = document.pop("_figure_images", {})
    for figure_id, image in figure_images.items():
        try:
            figure_dir = output / "figures"
            figure_dir.mkdir(parents=True, exist_ok=True)
            content_type = str(getattr(image, "content_type", "image/png"))
            extension = {"image/jpeg": ".jpg", "image/jpg": ".jpg", "image/gif": ".gif", "image/webp": ".webp"}.get(content_type, ".png")
            destination = figure_dir / f"{figure_id}{extension}"
            if hasattr(image, "save"):
                image.save(str(destination), format="PNG")
            else:
                blob = getattr(image, "blob", image)
                if not isinstance(blob, (bytes, bytearray)):
                    raise TypeError(f"unsupported figure payload: {type(blob).__name__}")
                destination.write_bytes(bytes(blob))
            for figure in document.get("figures", []):
                if figure.get("id") == figure_id:
                    figure["image_path"] = destination.relative_to(output).as_posix()
                    break
        except Exception as exc:
            document.setdefault("warnings", []).append(f"figure image write failed for {figure_id}: {type(exc).__name__}: {exc}")
    structure_path.write_text(json.dumps(document, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    write_markdown(document, markdown_path)
    table_paths: list[str] = []
    for table in document.get("tables", []):
        table_paths.append(write_table_csv(table, table_dir / f"{table['id']}.csv"))
    return {
        "structure_path": structure_path.relative_to(output).as_posix(),
        "content_path": markdown_path.relative_to(output).as_posix(),
        "table_paths": [Path(item).relative_to(output).as_posix() for item in table_paths],
    }


def manifest_document(document: dict[str, Any], path: Path, source_id: str, files: dict[str, str], path_base: Path) -> dict[str, Any]:
    has_content = any(document.get(key) for key in ("blocks", "sections", "tables", "figures", "references"))
    figure_image_paths = [item["image_path"] for item in document.get("figures", []) if item.get("image_path")]
    # A PDF that produced content after a backend failure, OCR mismatch, page
    # rendering issue, or other document-level warning is usable but not a
    # clean extraction. Keep that distinction visible to downstream review.
    status = "extracted" if has_content else ("partial" if document.get("warnings") else "empty")
    if document.get("kind") == "pdf" and has_content and document.get("warnings"):
        status = "partial"
    if status == "extracted" and not document.get("warnings"):
        review_status = "ready-for-llm"
        recommended_actions: list[str] = ["阅读 Markdown，并将最终事实绑定到原始来源锚点。"]
    elif status == "failed":
        review_status = "input-error"
        recommended_actions = ["检查文件身份、格式和环境报告后重试。"]
    elif status in {"empty", "partial"}:
        review_status = "review-required"
        recommended_actions = ["打开原始文件并抽查结构化产物。"]
        if document.get("kind") == "pdf":
            if any("OCR is disabled" in warning for warning in document.get("warnings", [])):
                recommended_actions.append("扫描 PDF 使用 --ocr，并保留 --render-pages 页图复核。")
            else:
                recommended_actions.append("检查双栏、表格、图注和页图；必要时切换 PDF backend。")
        if document.get("kind") == "docx" and document.get("figures"):
            recommended_actions.append("核对 DOCX 图像与正文/图注的对应关系。")
    else:
        review_status = "review-recommended"
        recommended_actions = ["阅读提取警告，再决定哪些片段可以进入语义推理。"]
    return {
        "source_id": source_id,
        "name": path.name,
        "path": relative_artifact_path(path, path_base),
        "sha256": sha256_file(path),
        "detected_kind": document["kind"],
        "extractor": document.get("extractor", document["kind"]),
        "backend_attempts": document.get("backend_attempts", []),
        "status": status,
        "review_status": review_status,
        "recommended_actions": recommended_actions,
        "content_signals": {
            "has_text": bool(document.get("blocks") or document.get("sections")),
            "has_tables": bool(document.get("tables")),
            "has_figures": bool(document.get("figures")),
            "has_references": bool(document.get("references")),
        },
        "page_count": document.get("page_count"),
        "section_count": len(document.get("sections", [])),
        "block_count": len(document.get("blocks", [])),
        "table_count": len(document.get("tables", [])),
        "figure_count": len(document.get("figures", [])),
        "figure_image_paths": figure_image_paths,
        "structure_path": files["structure_path"],
        "content_path": files["content_path"],
        "table_paths": files["table_paths"],
        "page_image_paths": document.get("page_image_paths", []),
        "warnings": document.get("warnings", []),
    }


def validate_bundle(bundle: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    for index, value in enumerate(bundle.get("root_input", [])):
        if is_absolute_reference(value):
            errors.append(f"root_input[{index}] must be relative to path_base")
    for index, item in enumerate(bundle.get("documents", [])):
        if not isinstance(item, dict):
            continue
        for key in ("path", "structure_path", "content_path", "table_paths", "page_image_paths", "figure_image_paths"):
            values = item.get(key, []) if key.endswith("_paths") else [item.get(key)]
            for value in values:
                if value and is_absolute_reference(value):
                    errors.append(f"documents[{index}].{key} must be relative to path_base")
    try:
        import jsonschema

        schema = json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))
        validator = jsonschema.validators.validator_for(schema)(schema)
        errors.extend(error.message for error in validator.iter_errors(bundle))
        return errors
    except ImportError:
        return errors


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("inputs", nargs="*", type=Path, help="PDF/XML/HTML/DOCX/XLSX/CSV/TSV/JSON/text files or directories")
    parser.add_argument("--output", type=Path, default=Path("source-extraction"), help="Output directory")
    parser.add_argument("--pdf-backend", choices=("auto", "docling", "pymupdf", "pdfplumber"), default="auto")
    ocr_group = parser.add_mutually_exclusive_group()
    ocr_group.add_argument("--ocr", dest="ocr", action="store_true", help="Opt in to Docling OCR for scanned or image-only PDFs")
    ocr_group.add_argument("--no-ocr", dest="ocr", action="store_false", help=argparse.SUPPRESS)
    parser.set_defaults(ocr=False)
    parser.add_argument("--extract-figures", action="store_true", help="Ask Docling to export detected figures as PNG assets")
    parser.add_argument("--render-pages", action="store_true", help="Render PDF pages to PNG for visual review")
    parser.add_argument("--render-dpi", type=int, default=144, help="DPI for --render-pages (72-300)")
    parser.add_argument("--model-cache", type=Path, default=None, help="Docling 模型缓存目录；默认使用 Docling cache/models")
    parser.add_argument("--no-model-download", action="store_true", help="不自动下载 Docling 模型；缺失时 auto PDF 允许回退")
    parser.add_argument("--force-model-download", action="store_true", help="强制重新下载当前 PDF profile 所需 Docling 模型")
    parser.add_argument("--check-environment", action="store_true", help="Print the active environment and exit")
    args = parser.parse_args()
    if args.check_environment:
        print(json.dumps(environment_report(args.model_cache, with_ocr=args.ocr), ensure_ascii=False, indent=2))
        return 0
    if not args.inputs:
        parser.error("provide at least one input path or use --check-environment")
    paths = iter_inputs(args.inputs, args.output)
    args.output.mkdir(parents=True, exist_ok=True)
    has_pdf = any(detected_kind(path) == "pdf" for path in paths)
    use_docling = has_pdf and args.pdf_backend in {"auto", "docling"}
    if use_docling:
        model_bootstrap = ensure_docling_models(
            args.model_cache,
            with_ocr=args.ocr,
            allow_download=not args.no_model_download,
            force=args.force_model_download,
        )
        extraction_model_cache = docling_model_cache(args.model_cache)
        print(f"DOCLING_MODELS {model_bootstrap.get('status')} cache={model_bootstrap.get('cache_dir')}")
    else:
        model_bootstrap = {
            "status": "not-needed",
            "download_requested": False,
            "download_performed": False,
            "reason": "no PDF input or an explicit non-Docling PDF backend was selected",
        }
        extraction_model_cache = None
    documents: list[dict[str, Any]] = []
    for index, path in enumerate(paths, start=1):
        source_id = f"IN-{index:04d}"
        try:
            document = extract_one(path, source_id, args.pdf_backend, not args.ocr, args.extract_figures, extraction_model_cache)
            if document["kind"] == "pdf" and args.render_pages:
                try:
                    document["page_image_paths"] = render_pdf_pages(path, source_id, args.output, args.render_dpi)
                except Exception as exc:
                    document.setdefault("warnings", []).append(f"page rendering failed: {type(exc).__name__}: {exc}")
            files = write_document_artifacts(document, args.output)
            documents.append(manifest_document(document, path, source_id, files, args.output))
            print(f"EXTRACTED {source_id} {path.name} via {document.get('extractor', document['kind'])}")
        except Exception as exc:
            documents.append({"source_id": source_id, "name": path.name, "path": relative_artifact_path(path, args.output), "sha256": sha256_file(path), "detected_kind": detected_kind(path), "extractor": "none", "status": "failed", "review_status": "input-error", "recommended_actions": ["检查文件身份、格式和环境报告后重试。"], "warnings": [f"unhandled extraction failure: {type(exc).__name__}: {exc}"]})
            print(f"FAILED {source_id} {path.name}: {exc}", file=sys.stderr)
    bundle = {
        "schema_version": "1.0",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "extractor_version": EXTRACTOR_VERSION,
        "path_base": ".",
        "root_input": [relative_artifact_path(path, args.output) for path in args.inputs],
        "environment": environment_report(args.model_cache, with_ocr=args.ocr),
        "policy": {
            "pdf_backend": args.pdf_backend,
            "ocr": args.ocr,
            "ocr_backend": "docling" if args.ocr and args.pdf_backend in {"auto", "docling"} else None,
            "extract_figures": args.extract_figures,
            "render_pages": args.render_pages,
            "model_download": model_bootstrap,
            "interpretation": "agent-required",
        },
        "documents": documents,
        "summary": {
            "file_count": len(documents),
            "extracted_count": sum(item.get("status") == "extracted" for item in documents),
            "partial_count": sum(item.get("status") == "partial" for item in documents),
            "empty_count": sum(item.get("status") == "empty" for item in documents),
            "failed_count": sum(item.get("status") == "failed" for item in documents),
            "review_needed_count": sum(item.get("review_status") in {"review-required", "review-recommended", "input-error"} for item in documents),
            "ready_count": sum(item.get("review_status") == "ready-for-llm" for item in documents),
            "table_count": sum(item.get("table_count", 0) for item in documents),
            "figure_count": sum(item.get("figure_count", 0) for item in documents),
        },
    }
    bundle_path = args.output / "source-extraction.json"
    bundle_path.write_text(json.dumps(bundle, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    errors = validate_bundle(bundle)
    if errors:
        (args.output / "validation-notes.txt").write_text("\n".join(errors) + "\n", encoding="utf-8")
        print(f"ERROR extraction bundle failed schema validation: {len(errors)} issue(s)", file=sys.stderr)
        return 1
    print(f"WROTE {bundle_path}; files={len(documents)}, tables={bundle['summary']['table_count']}, figures={bundle['summary']['figure_count']}")
    return 1 if documents and bundle["summary"]["failed_count"] == len(documents) else 0


if __name__ == "__main__":
    raise SystemExit(main())
