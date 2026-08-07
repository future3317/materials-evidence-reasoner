from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

import fitz

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from extract_sources import (  # noqa: E402
    detected_kind,
    docling_model_status,
    ensure_docling_models,
    environment_report,
    extract_delimited,
    extract_docx,
    extract_html,
    extract_json,
    extract_pdf,
    extract_pdf_pymupdf,
    extract_xlsx,
    extract_xml,
    manifest_document,
    write_document_artifacts,
)
from prepare_intake import detect_kind, iter_files  # noqa: E402
from validate_output import validate_relative_artifact_paths  # noqa: E402


JATS = """<?xml version="1.0" encoding="UTF-8"?>
<article xmlns:xlink="http://www.w3.org/1999/xlink">
  <front><article-meta>
    <article-title>JATS Fixture</article-title>
    <article-id pub-id-type="doi">10.0000/fixture</article-id>
    <abstract><p>Abstract with a traceable source anchor.</p></abstract>
  </article-meta></front>
  <body><sec id="methods"><title>Methods</title><p>Anneal at 300 C for 2 h.</p></sec>
  <sec id="results"><title>Results</title><p>The conductivity was 3.2 W/mK.</p>
    <table-wrap id="tbl1"><label>Table 1</label><caption><p>Conductivity</p></caption>
      <table><tr><th>Sample</th><th>Value</th></tr><tr><td>A</td><td>3.2</td></tr></table>
    </table-wrap>
  </sec></body>
  <back><ref-list><ref id="r1">Fixture reference.</ref></ref-list></back>
</article>"""


class SourceExtractionTests(unittest.TestCase):
    def test_xml_preserves_structure_and_anchors(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "fixture.xml"
            path.write_text(JATS, encoding="utf-8")
            document = extract_xml(path, "IN-0001")
        self.assertEqual(document["title"], "JATS Fixture")
        self.assertEqual(document["metadata"]["doi"], "10.0000/fixture")
        self.assertTrue(any(section["anchor"].startswith("IN-0001:section=methods") for section in document["sections"]))
        self.assertEqual(document["tables"][0]["headers"], ["Sample", "Value"])
        self.assertEqual(document["references"][0]["id"], "r1")

    def test_pdf_fallback_keeps_page_anchors(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "fixture.pdf"
            pdf = fitz.open()
            page = pdf.new_page()
            page.insert_text((72, 72), "PDF fixture conductivity 3.2 W/mK")
            pdf.save(path)
            pdf.close()
            document = extract_pdf_pymupdf(path, "IN-0002")
        self.assertEqual(document["extractor"], "pymupdf")
        self.assertEqual(document["page_count"], 1)
        self.assertIn("page=1", document["blocks"][0]["anchor"])
        self.assertIn("conductivity", document["blocks"][0]["text"])

    def test_environment_report_exposes_capabilities_and_backend_attempts(self) -> None:
        report = environment_report()
        self.assertIn("capabilities", report)
        self.assertIn("pdf_text", report["capabilities"])
        self.assertIn("install_hints", report)
        self.assertIn("profile_status", report)
        self.assertIn("recommended_actions", report)
        self.assertIn("model_cache", report)
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "fixture.pdf"
            pdf = fitz.open()
            pdf.new_page()
            pdf.save(path)
            pdf.close()
            document = extract_pdf(path, "IN-0007", "pymupdf")
        self.assertEqual(document["backend_attempts"], [{"backend": "pymupdf", "status": "used"}])

    def test_empty_pdf_explains_ocr_next_step(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "scan-like.pdf"
            pdf = fitz.open()
            pdf.new_page()
            pdf.save(path)
            pdf.close()
            document = extract_pdf(path, "IN-0008", "pymupdf")
        self.assertEqual(sum("OCR is disabled" in warning for warning in document["warnings"]), 1)

    def test_ocr_request_is_not_silent_on_text_backend(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "scan-like.pdf"
            pdf = fitz.open()
            pdf.new_page()
            pdf.save(path)
            pdf.close()
            document = extract_pdf(path, "IN-0009", "pymupdf", no_ocr=False)
        self.assertTrue(any("does not perform OCR" in warning for warning in document["warnings"]))

    def test_pdf_with_document_warning_is_marked_partial(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "fallback.pdf"
            path.write_bytes(b"fixture")
            manifest = manifest_document(
                {
                    "kind": "pdf",
                    "blocks": [{"text": "usable text"}],
                    "sections": [],
                    "tables": [],
                    "figures": [],
                    "references": [],
                    "warnings": ["docling failed; used fallback"],
                },
                path,
                "IN-0011",
                {"structure_path": "documents/IN-0011.json", "content_path": "documents/IN-0011.md", "table_paths": []},
                path.parent,
            )
        self.assertEqual(manifest["status"], "partial")
        self.assertEqual(manifest["review_status"], "review-required")
        self.assertTrue(manifest["recommended_actions"])

    def test_html_and_intake_kind_detection(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "fixture.html"
            path.write_text("<html><title>Fixture</title><h1>Results</h1><p>Value: 3.2</p></html>", encoding="utf-8")
            document = extract_html(path, "IN-0003")
            self.assertEqual(document["title"], "Fixture")
            self.assertTrue(document["sections"])
            xml_path = Path(directory) / "paper.xml"
            xml_path.write_text("<article/>", encoding="utf-8")
            docx_path = Path(directory) / "paper.docx"
            docx_path.write_bytes(b"PK\x03\x04")
            self.assertEqual(detect_kind(xml_path)[0], "xml")
            self.assertEqual(detect_kind(docx_path)[0], "docx")

    def test_cli_writes_relative_bundle_paths(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            xml = root / "fixture.xml"
            xml.write_text(JATS, encoding="utf-8")
            output = root / "out"
            result = subprocess.run(
                [sys.executable, str(SCRIPTS / "extract_sources.py"), str(xml), "--output", str(output)],
                cwd=ROOT,
                text=True,
                capture_output=True,
                encoding="utf-8",
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            bundle = json.loads((output / "source-extraction.json").read_text(encoding="utf-8"))
            item = bundle["documents"][0]
            self.assertEqual(bundle["path_base"], ".")
            self.assertFalse(Path(bundle["root_input"][0]).is_absolute())
            self.assertFalse(Path(item["path"]).is_absolute())
            self.assertEqual(item["status"], "extracted")
            self.assertEqual(bundle["policy"]["model_download"]["status"], "not-needed")
            self.assertFalse(Path(item["content_path"]).is_absolute())
            self.assertTrue((output / item["content_path"]).is_file())
            self.assertTrue((output / item["structure_path"]).is_file())
            dashboard = root / "source-dashboard.html"
            dashboard_result = subprocess.run(
                [sys.executable, str(SCRIPTS / "render_source_dashboard.py"), str(output / "source-extraction.json"), "-o", str(dashboard)],
                cwd=ROOT,
                text=True,
                capture_output=True,
                encoding="utf-8",
            )
            self.assertEqual(dashboard_result.returncode, 0, dashboard_result.stderr)
            dashboard_text = dashboard.read_text(encoding="utf-8")
            self.assertIn("文献提取检查台", dashboard_text)
            self.assertIn("Docling 模型", dashboard_text)
            self.assertIn("原始文件", dashboard_text)
            self.assertNotIn("fetch(", dashboard_text)

    def test_csv_preserves_newlines_inside_quoted_fields(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "multiline.csv"
            path.write_text('Sample,Note\nA,"line one\nline two"\n', encoding="utf-8")
            document = extract_delimited(path, "IN-0010")
        self.assertEqual(document["tables"][0]["rows"], [["A", "line one\nline two"]])

    def test_intake_does_not_reingest_its_output_subtree(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            output = root / "intake-output"
            output.mkdir()
            (root / "measurements.csv").write_text("Sample,Value\nA,1\n", encoding="utf-8")
            (output / "old-packet.json").write_text("{}", encoding="utf-8")
            files = iter_files(root, output)
        self.assertEqual([path.name for path in files], ["measurements.csv"])

    def test_artifact_path_validator_rejects_uri_schemes(self) -> None:
        errors = validate_relative_artifact_paths(
            {
                "artifact_manifest": [{"path": "javascript:alert(1)"}],
                "input_assessment": {"received_inputs": [{"path": "C:relative.txt"}]},
            }
        )
        self.assertEqual(len(errors), 2)

    def test_xlsx_preserves_formula_and_cached_value_metadata(self) -> None:
        import openpyxl

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "fixture.xlsx"
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "Results"
            sheet.append(["Sample", "Value"])
            sheet.append(["A", "=1+2"])
            workbook.save(path)
            workbook.close()
            document = extract_xlsx(path, "IN-0004")
        cell = document["tables"][0]["rows"][0][1]
        self.assertEqual(cell["formula"], "=1+2")
        self.assertIsNone(cell["cached_value"])
        self.assertTrue(any("formula cells" in warning for warning in document["tables"][0]["warnings"]))

    def test_magic_sniff_overrides_wrong_extension(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "paper.bin"
            pdf = fitz.open()
            pdf.new_page()
            pdf.save(path)
            pdf.close()
            self.assertEqual(detected_kind(path), "pdf")

    def test_supplementary_csv_and_json_are_tabulated(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            csv_path = root / "supplement.csv"
            csv_path.write_text("Sample,Value\nA,3.2\n", encoding="utf-8")
            json_path = root / "supplement.json"
            json_path.write_text(json.dumps({"measurements": [{"sample": "A", "value": 3.2}]}), encoding="utf-8")
            from extract_sources import extract_delimited, extract_json

            csv_document = extract_delimited(csv_path, "IN-0005")
            json_document = extract_json(json_path, "IN-0006")
        self.assertEqual(csv_document["tables"][0]["rows"], [["A", "3.2"]])
        self.assertEqual(json_document["tables"][0]["headers"], ["sample", "value"])

    def test_json_extraction_preserves_multiple_object_arrays(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "supplement.json"
            path.write_text(json.dumps({"measurements": [{"sample": "A", "value": 3.2}], "conditions": [{"sample": "A", "temperature": 300}]}), encoding="utf-8")
            document = extract_json(path, "IN-0012")
        self.assertEqual([table["id"] for table in document["tables"]], ["IN-0012-T001", "IN-0012-T002"])
        self.assertEqual(document["metadata"]["arrays"], ["measurements", "conditions"])
        self.assertEqual(len(document["reading_order"]), 2)

    def test_docx_figure_extraction_is_opt_in(self) -> None:
        import base64
        from docx import Document

        png = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII=")
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            image = root / "figure.png"
            image.write_bytes(png)
            path = root / "fixture.docx"
            doc = Document()
            doc.add_heading("Results", level=1)
            doc.add_paragraph("A result paragraph.")
            doc.add_picture(str(image))
            doc.save(path)
            extracted = extract_docx(path, "IN-0013", extract_figures=True)
            output = root / "out"
            files = write_document_artifacts(extracted, output)
            figure_written = bool(list((output / "figures").glob("IN-0013-F001.*")))
            markdown_written = (output / files["content_path"]).is_file()
        self.assertTrue(extracted["figures"])
        self.assertIn("IN-0013-F001", extracted.get("figures", [{}])[0].get("id", ""))
        self.assertTrue(any(item["type"] == "figure" for item in extracted["reading_order"]))
        self.assertTrue(figure_written)
        self.assertTrue(markdown_written)

    def test_docling_model_status_does_not_download(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            status = docling_model_status(Path(directory) / "empty-model-cache")
        self.assertIn(status["status"], {"missing", "unavailable"})
        self.assertFalse(status.get("download_performed", False))

    def test_docling_model_download_can_be_disabled(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            status = docling_model_status(Path(directory) / "empty-model-cache")
            result = ensure_docling_models(Path(directory) / "empty-model-cache", allow_download=False)
        if status["status"] == "missing":
            self.assertEqual(result["status"], "skipped")
            self.assertIn("--no-model-download", result["next_step"])
        else:
            self.assertEqual(result["status"], "unavailable")


if __name__ == "__main__":
    unittest.main()
